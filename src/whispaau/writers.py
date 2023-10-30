import csv
import json
from datetime import datetime
from typing import TextIO

import docx
from whisper.utils import ResultWriter, format_timestamp


class WriteCSV(ResultWriter):
    extension: str = "csv"

    def write_result(self, result: dict, file: TextIO, options: dict):
        fieldnames: list[str] = list(result["segments"][0].keys())
        writer = csv.DictWriter(file, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(result["segments"])


class WriteDOTE(ResultWriter):
    extension: str = "dote.json"

    @staticmethod
    def format_result(result: dict):
        interface = {"lines": []}
        for line in result["segments"]:
            line_add = {
                "startTime": format_timestamp(line["start"], True),
                "endTime": format_timestamp(line["end"], True),
                "speakerDesignation": "",
                "text": line["text"].strip(),
            }
            interface["lines"].append(line_add)
        return interface

    def write_result(self, result: dict, file: TextIO, options: dict):
        result = self.format_result(result)
        json.dump(result, file)


class WriteDOCX(ResultWriter):
    extension: str = "docx"

    @staticmethod
    def format_time(
        start_time: int | float, end_time: int | float, max_time: int | float
    ) -> str:
        dt_start = datetime.utcfromtimestamp(start_time)
        dt_end = datetime.utcfromtimestamp(end_time)
        time_format = "%M:%S"
        if max_time >= 3600:
            time_format = "%H:%M:%S"
        return f"{dt_start.strftime(time_format)} - {dt_end.strftime(time_format)}"

    def write_result(self, result: dict, file: TextIO, options: dict):
        audio_filename = options.get("filename", "").stem
        document = docx.Document()
        document.add_heading(audio_filename, level=2)
        document.add_heading(options["jobname"], level=4)
        document.extended_properties.set_property("total_time", "   1")
        document.extended_properties.set_property(
            "application", "Whisper Transcription AAU extension"
        )
        document.core_properties.title = audio_filename
        document.core_properties.author = options.get("username", "")
        document.core_properties.subject = options.get("jobname", "")

        p = document.add_paragraph()
        max_time = result["segments"][-1]["end"]
        for line in result["segments"]:
            time = p.add_run(self.format_time(line["start"], line["end"], max_time))
            time.italic = True
            p.add_run("\t")
            p.add_run(f'{line["text"].strip()}\n')
        document.save(file.name)
